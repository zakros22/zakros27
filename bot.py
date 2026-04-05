import os
import telebot
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
import tempfile
import re
from groq import Groq
import PyPDF2
import docx
from datetime import datetime

# ================== الإعدادات ==================
BOT_TOKEN = os.environ.get("BOT_TOKEN")
GROQ_API_KEY = os.environ.get("GROQ_API_KEY")

if not BOT_TOKEN or not GROQ_API_KEY:
    raise Exception("❌ يرجى تعيين BOT_TOKEN و GROQ_API_KEY في متغيرات البيئة")

bot = telebot.TeleBot(BOT_TOKEN)
client = Groq(api_key=GROQ_API_KEY)

# اللهجات المدعومة
DIALECTS = {
    "iraqi": "🇮🇶 اللهجة العراقية",
    "syrian": "🇸🇾 اللهجة السورية",
    "fusha": "📖 الفصحى",
    "egyptian": "🇪🇬 اللهجة المصرية",
    "gulf": "🇦🇪 اللهجة الخليجية",
    "moroccan": "🇲🇦 اللهجة المغربية"
}

# تخزين مؤقت لبيانات المستخدم
user_data = {}

# ================== دوال استخراج النص من الملفات ==================
def extract_text_from_file(file_path):
    """استخراج النص من ملف (txt, pdf, docx)"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif ext == '.pdf':
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
    elif ext == '.docx':
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return None

def create_file_with_translated_text(original_path, translated_text, output_path):
    """إنشاء ملف جديد بنفس نوع الأصلي مع النص المترجم"""
    ext = os.path.splitext(original_path)[1].lower()
    if ext == '.txt':
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(translated_text)
    elif ext == '.pdf':
        # PDF معقد: سنقوم بإنشاء ملف PDF جديد بسيط يحتوي على النص المترجم
        # (للبساطة، سنستخدم تقرير PDF بسيط عبر مكتبة reportlab)
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4
        y = height - 40
        for line in translated_text.split('\n'):
            if y < 40:
                c.showPage()
                y = height - 40
            c.drawString(40, y, line[:100])
            y -= 15
        c.save()
    elif ext == '.docx':
        from docx import Document
        doc = Document()
        doc.add_paragraph(translated_text)
        doc.save(output_path)
    else:
        # لأي نوع آخر، نحفظ كنص عادي
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(translated_text)

# ================== الترجمة باستخدام Groq ==================
def translate_text(text, dialect_name):
    """ترجمة النص إلى اللهجة المطلوبة"""
    prompt = f"""أنت مترجم خبير. حول النص التالي إلى {dialect_name} بدقة مع الحفاظ على المعنى الأصلي.
أخرج النص المترجم فقط، بدون أي إضافات أو تعليقات أو علامات تنصيص.

النص:
{text}"""
    try:
        completion = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=2000
        )
        return completion.choices[0].message.content.strip()
    except Exception as e:
        raise Exception(f"Groq API error: {e}")

# ================== أوامر البوت ==================
@bot.message_handler(commands=['start'])
def start(message):
    bot.send_message(
        message.chat.id,
        "🌍 *بوت ترجمة الملفات إلى اللهجات العربية*\n\n"
        "📤 أرسل لي ملفاً (txt, pdf, docx) وسأقوم بترجمته إلى اللهجة التي تختارها.\n"
        "📏 الحد الأقصى للنص: 2500 حرف.\n"
        "⚠️ الملفات غير النصية (صور، فيديو، صوت) غير مدعومة حالياً.\n\n"
        "👈 أرسل الملف الآن.",
        parse_mode="Markdown"
    )

@bot.message_handler(content_types=['document'])
def handle_document(message):
    file_name = message.document.file_name
    file_ext = os.path.splitext(file_name)[1].lower()
    supported = ['.txt', '.pdf', '.docx']
    if file_ext not in supported:
        bot.reply_to(message, f"❌ نوع الملف غير مدعوم. الأنواع المدعومة: {', '.join(supported)}")
        return

    try:
        # تحميل الملف
        file_info = bot.get_file(message.document.file_id)
        downloaded = bot.download_file(file_info.file_path)
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
            tmp.write(downloaded)
            tmp_path = tmp.name

        # استخراج النص
        text = extract_text_from_file(tmp_path)
        if not text or len(text.strip()) < 10:
            bot.reply_to(message, "❌ لم يتم العثور على نص صالح في هذا الملف.")
            os.unlink(tmp_path)
            return
        if len(text) > 2500:
            bot.reply_to(message, "❌ النص في الملف طويل جداً (الحد الأقصى 2500 حرف).")
            os.unlink(tmp_path)
            return

        # حفظ بيانات الجلسة
        user_data[message.chat.id] = {
            "original_path": tmp_path,
            "original_name": file_name,
            "original_text": text
        }

        # عرض أزرار اللهجات
        markup = InlineKeyboardMarkup()
        for key, name in DIALECTS.items():
            markup.add(InlineKeyboardButton(name, callback_data=key))
        bot.send_message(message.chat.id, "✨ اختر اللهجة التي تريد الترجمة إليها:", reply_markup=markup)

    except Exception as e:
        bot.reply_to(message, f"❌ فشل تحميل الملف: {str(e)[:100]}")
        if 'tmp_path' in locals():
            os.unlink(tmp_path)

@bot.callback_query_handler(func=lambda call: call.data in DIALECTS)
def handle_dialect(call):
    user_id = call.message.chat.id
    dialect_key = call.data
    dialect_name = DIALECTS[dialect_key]
    session = user_data.get(user_id)
    if not session:
        bot.answer_callback_query(call.id, "انتهت صلاحية الجلسة، أعد إرسال الملف.", show_alert=True)
        return

    original_text = session["original_text"]
    original_path = session["original_path"]
    original_name = session["original_name"]

    bot.answer_callback_query(call.id, f"جاري الترجمة إلى {dialect_name}...")
    msg = bot.send_message(user_id, f"⏳ جاري ترجمة النص إلى {dialect_name}... قد تستغرق 10-30 ثانية.")

    try:
        translated = translate_text(original_text, dialect_name)
        # إنشاء ملف جديد مترجم
        base, ext = os.path.splitext(original_name)
        new_name = f"{base}_{dialect_key}{ext}"
        new_path = tempfile.mktemp(suffix=ext)
        create_file_with_translated_text(original_path, translated, new_path)

        with open(new_path, 'rb') as f:
            bot.send_document(user_id, f, caption=f"✅ تمت الترجمة إلى {dialect_name}", visible_file_name=new_name)

        # تنظيف الملفات
        os.unlink(original_path)
        os.unlink(new_path)
        bot.delete_message(user_id, msg.message_id)
        del user_data[user_id]

    except Exception as e:
        bot.edit_message_text(f"❌ فشلت الترجمة: {str(e)[:200]}", user_id, msg.message_id)
        # تنظيف
        if os.path.exists(original_path):
            os.unlink(original_path)
        if user_id in user_data:
            del user_data[user_id]

if __name__ == "__main__":
    print("✅ بوت الترجمة إلى اللهجات العربية يعمل...")
    bot.remove_webhook()
    bot.infinity_polling()
